from docxPdfImage import *
from docx.enum.text import WD_COLOR_INDEX
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import re
import docx
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

"""
Testing iter_block_items()
"""
def iter_block_items(parent):
    """
    Tạo tham chiếu đến từng đoạn và bảng con trong file doc, theo thứ tự tài liệu. 
    Mỗi giá trị trả về là một thể hiện của Bảng hoặc Đoạn văn. 
    'parent' thường là một tham chiếu đến một chính Đối tượng tài liệu, 
    hoạt động cho đối tượng _Cell | đoạn văn | bảng
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def color_string(key,countKey,p1,p):
    ##    tô vàng key
##    input: key, số thứ tự key, đoạn văn chứa key, đoạn mới chứa key được tô vàng
##    output: đoạn văn đã được tô vàng, số thứ tự key
    str1 =""
    sub_end =""
    #print(p1,"\n")
    for i in range(len(key)):
        substrings = p1.split(key[i]) # split đoạn
        if(i!=len(key)-1):
            p1 = str1.join(substrings)
    #print(substrings)  
        for substring in substrings[:-1]:
            #print('subs',substring)
            countKey += 1
            p.add_run(substring,style = 'CommentsStyle') # Ép kiểu chữ theo font'CommentStyle'
            font = p.add_run(key[i],style = 'CommentsStyle').font.highlight_color = WD_COLOR_INDEX.YELLOW # tô vàng key
            count = str(countKey)
            font = p.add_run(count,style = 'CommentsStyle').font.highlight_color = WD_COLOR_INDEX.RED # tô đỏ số thứ tự của key
        sub_end = substrings[-1]
    p.add_run(substrings[-1], style = 'CommentsStyle')
    return countKey
def Size(filename):# tìm và chọn size của văn bản
    size = []
    doc = docx.Document(filename)
    for p in doc.paragraphs:
        for i in p.runs:
            if i.font.size != None:
                size.append(i.font.size/12700)
    return size
def iter_unique_cells(row): #(Hợp nhất cells theo dòng, bỏ qua các lần lặp lại)
    prior_tc = None
    for cell in row.cells:
        this_tc = cell._tc
        if this_tc is prior_tc:
            continue
        prior_tc = this_tc
        yield cell
def findColor(filename,key,newName):
    ##    tìm và tô vàng key
##    input: file cần xử ký, key cần tìm và tô vàng
##    output: file đã tô vàng và đánh thứ tự cho key
    countKey = 0 # khởi tạo số thứ tự key
    doc = Document(filename)
    
    #Tạo font theo 'CommentsStyle'
    par = doc.paragraphs[0]
    font_styles = doc.styles
    font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    if '{par.style.font.name}'!= None:
        font_object.name = f"'{par.style.font.name}'"
    sizes = Size(filename)
    run = 0
    for p in doc.paragraphs:
        for i in p.runs:
            if i.font.size != None:
                font_object.size = Pt(sizes[run])
                run += 1
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            #print(type(block)
            p1 = block.text
            match = re.findall(key,p1,re.IGNORECASE)
            #print(match)
            #for igkey in matc
            if len(match)>0: #so khớp không phân biệt hoa thường
                block.text = ""
                countKey=color_string(match,countKey,p1,block)
        else:
            for row in block.rows:
                for p in iter_unique_cells(row):
                    p1 = p.text
                    match = re.findall(key,p1,re.IGNORECASE)
                    if len(match)>0: #so khớp không phân biệt hoa thường
                        p.text = " "
                        p = p.add_paragraph()
                        print(p)
                        countKey = color_string(match,countKey,p1,p)
    doc.save(newName)
    return countKey,key

def replace_string(key,value,NumberList,countKey,p):
    lenght = len(key)# độ dài của key
    tmp_padding = len(key) - len (value)# so sánh độ dài của key và value
    matchs = re.findall(key,p.text,re.IGNORECASE) #khớp không phân biệt
    lines = p.runs 
    for j in range(len(lines)):# cho j vào trong các dòng
        padding = 0 #khởi tạo padding
        line = lines[j].text 
        for i in range(len(line)-lenght+1): #tạo i là khoảng thay thế của value
            text = line[i - padding : i + lenght - padding] #thay thế value vào vị trí của key khớp khoảng cách khi thay đổi
            if text in matchs:
                if countKey in NumberList:
                    text = line.replace(text, value) # bắt đầu thay đổi(hàm replace giữ lại font chứ ban đầu)
                    padding -= tmp_padding # thay thế khoảng cách của từ
                    lines[j].text = text #dòng text chứa value
                countKey +=1 
    return countKey

def replace(filename,key,value,numberList,output_file):
##    hàm duyệt từng đoạn trong file
##    tìm và thay thế từ ở vị trí chỉ định
##    input: tên file, từ muốn đổi, từ để đổi, danh sách vị trí đổi
##    output: file word đã được thay từ ở những vị trí chỉ định
    countKey = 1 # khởi tạo số thứ tự key
    doc = Document(filename)
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            if re.findall(key,block.text,re.IGNORECASE): #so khớp không phân biệt hoa thường
                #print(re.findall(key,block.text,re.IGNORECASE))
                countKey = replace_string(key,value,numberList,countKey,block)
        else:
                for row in block.rows:
                    for cell in iter_unique_cells(row):
                        for p in cell.paragraphs:
                            if re.findall(key,p.text,re.IGNORECASE): #so khớp không phân biệt hoa thường
                                 #print(re.search(key,p.text,re.IGNORECASE))
                                 countKey = replace_string(key,value,numberList,countKey,p)
    doc.save(output_file)
'''input_file = 'output/phong8.docx'
input_file = os.path.abspath(input_file)
#file = os.getcwd() + "/" + input_file
key = u'công việc'
value = u'công việc tuần'
numberList=[1,2]
output_file = 'output/output.docx'
replace(input_file,key,value,numberList,output_file)'''